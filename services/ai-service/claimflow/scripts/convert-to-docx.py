import os
import re
from docx import Document
from docx.shared import Inches

def convert_md_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    doc = Document()
    
    with open(md_path, 'r') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        
        # Images: ![caption](/path/to/image.png)
        elif line.startswith('!['):
            match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
            if match:
                caption, img_path = match.groups()
                if os.path.exists(img_path):
                    doc.add_picture(img_path, width=Inches(6))
                    doc.add_paragraph(caption, style='Caption')
                else:
                    print(f"Warning: Image {img_path} not found.")
        
        # Alerts (simplified)
        elif line.startswith('> [!'):
            # Just add as a noted paragraph
            text = re.sub(r'^> \[!.*?\]', '', line).strip()
            p = doc.add_paragraph()
            p.add_run("NOTE: ").bold = True
            p.add_run(text)

        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.', line):
            doc.add_paragraph(re.sub(r'^\d+\.\s*', '', line), style='List Number')

        # Tables (simplified extraction of text)
        elif '|' in line and not line.startswith('|-'):
            # Skip separator lines
            if re.match(r'^[|\-\s:]+$', line):
                continue
            cells = [c.strip() for c in line.split('|') if c.strip()]
            if cells:
                doc.add_paragraph(" | ".join(cells))
        
        # Normal text
        else:
            doc.add_paragraph(line)

    doc.save(docx_path)
    print(f"Successfully converted {md_path} to {docx_path}")

if __name__ == "__main__":
    report_md = "/home/pipi-ngure/.gemini/antigravity/brain/e725d2c5-2d57-4611-aa07-055a5455c4b6/performance_report.md"
    report_docx = "/home/pipi-ngure/.gemini/antigravity/brain/e725d2c5-2d57-4611-aa07-055a5455c4b6/performance_report.docx"
    convert_md_to_docx(report_md, report_docx)
