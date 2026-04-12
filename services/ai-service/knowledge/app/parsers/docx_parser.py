"""
DOCX document parser using python-docx.
Extracts text with pseudo-page structure and heading detection.
"""

from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from structlog import get_logger

from app.parsers.base import ParsedDocument, ParsedPage

logger = get_logger(__name__)

_LINES_PER_PAGE = 45


def parse_docx(file_path: str | Path) -> ParsedDocument:
    """
    Parse a DOCX file and extract text with structure.
    DOCX doesn't have real pages, so we estimate based on paragraph count.

    @param file_path: Path to the DOCX file
    @returns ParsedDocument with pseudo-pages and headings
    """
    doc = DocxDocument(str(file_path))
    metadata: dict[str, str] = {}

    core = doc.core_properties
    if core.title:
        metadata["title"] = core.title
    if core.author:
        metadata["author"] = core.author
    if core.subject:
        metadata["subject"] = core.subject

    current_page_lines: list[str] = []
    current_sections: list[str] = []
    pages: list[ParsedPage] = []
    page_num = 1

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            current_page_lines.append("")
            continue

        style_name = (para.style.name or "").lower() if para.style else ""
        is_heading = "heading" in style_name or (
            para.runs
            and para.runs[0].bold
            and len(text) < 100
            and not text.endswith(".")
        )

        if is_heading:
            current_sections.append(text)

        current_page_lines.append(text)

        if len(current_page_lines) >= _LINES_PER_PAGE:
            pages.append(
                ParsedPage(
                    page_number=page_num,
                    text="\n".join(current_page_lines),
                    section_titles=current_sections.copy(),
                )
            )
            page_num += 1
            current_page_lines = []
            current_sections = []

    if current_page_lines:
        pages.append(
            ParsedPage(
                page_number=page_num,
                text="\n".join(current_page_lines),
                section_titles=current_sections,
            )
        )

    logger.info("docx_parsed", file=str(file_path), pages=len(pages))
    return ParsedDocument(pages=pages, total_pages=len(pages), metadata=metadata)
